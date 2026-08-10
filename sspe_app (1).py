import streamlit as st
import anthropic
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="SSPE – Sports Science Publication Engine",
    page_icon="🏅",
    layout="wide",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .main { background-color: #0e1117; }
    .stButton>button {
        background-color: #1f6feb;
        color: white;
        border-radius: 8px;
        font-weight: bold;
        width: 100%;
    }
    .stButton>button:hover { background-color: #388bfd; }
    .title-box {
        background: linear-gradient(135deg, #1f6feb 0%, #0d419d 100%);
        padding: 2rem;
        border-radius: 12px;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-card {
        background-color: #161b22;
        border: 1px solid #30363d;
        border-radius: 10px;
        padding: 1.5rem;
        margin-bottom: 1rem;
    }
    .output-box {
        background-color: #0d1117;
        border: 1px solid #21262d;
        border-radius: 8px;
        padding: 1.5rem;
        font-family: 'Georgia', serif;
        line-height: 1.8;
    }
</style>
""", unsafe_allow_html=True)

# ── Header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="title-box">
    <h1 style="color:white; margin:0;">🏅 Sports Science Publication Engine</h1>
    <p style="color:#a0aec0; margin-top:0.5rem;">Powered by Claude AI · Scopus & Sinta Ready</p>
</div>
""", unsafe_allow_html=True)

# ── API Key input ─────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Konfigurasi API")
    api_key = st.text_input("Anthropic API Key", type="password", placeholder="sk-ant-...")
    st.markdown("---")
    st.markdown("### 📖 Panduan")
    st.markdown("""
    1. Masukkan API Key
    2. Isi form konfigurasi
    3. Isi detail penelitian
    4. Klik **Generate**
    5. Download hasil (.docx)
    """)
    st.markdown("---")
    st.markdown("*SSPE v1.0 · Prof Bayu*")

# ── Main form ─────────────────────────────────────────────────────────────────
st.markdown("## 📋 Konfigurasi Artikel Ilmiah")

col1, col2 = st.columns(2)

with col1:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown("### [1] Jenis Penelitian")
    jenis = st.radio("", [
        "A. Kuantitatif / Eksperimen (SPSS, ANOVA, Regresi, Tes Fisik)",
        "B. Kualitatif (Studi Kasus, Etnografi, Fenomenologi)",
        "C. Systematic Literature Review / Meta-Analisis (PRISMA)"
    ], key="jenis")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown("### [3] Sumber Data")
    sumber = st.radio("", [
        "A. Upload File (Excel/CSV + Referensi)",
        "B. Input Manual / Ringkasan Data"
    ], key="sumber")
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown("### [2] Target Jurnal")
    target = st.radio("", [
        "A. Scopus Q1 / Q2",
        "B. Scopus Q3 / Q4",
        "C. Sinta 1 / 2",
        "D. Sinta 3 / 4 / 5 / 6"
    ], key="target")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown("### [4] Bahasa Output")
    bahasa = st.radio("", [
        "Bahasa Indonesia",
        "English (Academic)",
        "Bilingual (Indonesia + English)"
    ], key="bahasa")
    st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.markdown("### [5] Bagian yang Digenerate")
bagian = st.multiselect("Pilih bagian (bisa lebih dari satu):", [
    "Full Paper (A-Z) – Modular",
    "Judul & Abstrak",
    "Pendahuluan & Gap Analysis (SOTA)",
    "Metodologi Penelitian",
    "Hasil & Pembahasan Kritis",
    "Kesimpulan & Rekomendasi",
    "Daftar Pustaka (APA 7th)"
], default=["Full Paper (A-Z) – Modular"])
st.markdown('</div>', unsafe_allow_html=True)

# ── Research detail input ─────────────────────────────────────────────────────
st.markdown("## ✏️ Detail Penelitian")

st.markdown('<div class="section-card">', unsafe_allow_html=True)
col3, col4 = st.columns(2)
with col3:
    topik = st.text_input("🎯 Topik / Judul Penelitian", placeholder="Contoh: Pengaruh HIIT terhadap VO2Max atlet renang...")
    variabel_bebas = st.text_input("📊 Variabel Bebas / Independen", placeholder="Contoh: Latihan HIIT, Frekuensi latihan")
    variabel_terikat = st.text_input("📈 Variabel Terikat / Dependen", placeholder="Contoh: VO2Max, Daya tahan kardiovaskular")
with col4:
    subjek = st.text_input("👥 Subjek / Populasi Penelitian", placeholder="Contoh: 30 atlet renang usia 15-18 tahun")
    metode_analisis = st.text_input("🔬 Metode Analisis", placeholder="Contoh: Paired T-Test, ANOVA, Regresi Linear")
    durasi = st.text_input("📅 Durasi / Periode Penelitian", placeholder="Contoh: 8 minggu, pre-post test design")

ringkasan_data = st.text_area("📋 Ringkasan Data / Hasil Temuan", 
    placeholder="Contoh: Kelompok eksperimen mengalami peningkatan VO2Max rata-rata 12.3% (pre: 42.1 ml/kg/min → post: 47.3 ml/kg/min), p<0.05...",
    height=120)

referensi_kunci = st.text_area("📚 Referensi Kunci (opsional)", 
    placeholder="Contoh:\n- Smith et al. (2022). HIIT and VO2Max. Journal of Sports Science...\n- Johnson (2021). Swimming performance...",
    height=100)
st.markdown('</div>', unsafe_allow_html=True)

# ── Upload file section ───────────────────────────────────────────────────────
if "A. Upload File" in sumber:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown("### 📁 Upload File Pendukung")
    uploaded_file = st.file_uploader("Upload data (CSV/Excel)", type=["csv", "xlsx", "xls"])
    if uploaded_file:
        st.success(f"✅ File '{uploaded_file.name}' berhasil diupload")
    st.markdown('</div>', unsafe_allow_html=True)

# ── Helper functions ──────────────────────────────────────────────────────────
def build_prompt(jenis, target, bagian, bahasa, topik, variabel_bebas, 
                 variabel_terikat, subjek, metode_analisis, durasi, 
                 ringkasan_data, referensi_kunci):
    
    lang_instruction = {
        "Bahasa Indonesia": "Tulis seluruh artikel dalam Bahasa Indonesia akademik yang formal dan baku.",
        "English (Academic)": "Write the entire article in high-level academic English suitable for international journals.",
        "Bilingual (Indonesia + English)": "Write each section in both Bahasa Indonesia and English. Start with Indonesian, then provide the English version below each section."
    }.get(bahasa, "Bahasa Indonesia")

    quality_level = {
        "A. Scopus Q1 / Q2": "highest international standard (Nature, Lancet level). Use advanced statistical interpretation, critical synthesis of literature, theoretical framework, and implications for global sport science.",
        "B. Scopus Q3 / Q4": "solid international standard. Include clear methodology, statistical analysis, and discussion linked to existing literature.",
        "C. Sinta 1 / 2": "high national accreditation standard. Structured IMRAD format, clear analysis, and relevant discussion for Indonesian sport science context.",
        "D. Sinta 3 / 4 / 5 / 6": "standard national journal format. Clear and systematic writing with basic statistical analysis."
    }.get(target, "standard")

    sections_to_generate = ", ".join(bagian) if bagian else "Full Paper"

    prompt = f"""You are a Senior Sports Science Professor and academic writing expert with 20+ years of experience publishing in top-tier journals.

TASK: Generate a high-quality academic article based on the specifications below.

═══════════════════════════════════════
ARTICLE SPECIFICATIONS:
═══════════════════════════════════════
Research Type: {jenis}
Target Journal: {target} — Write at {quality_level}
Language: {lang_instruction}
Sections to Generate: {sections_to_generate}

═══════════════════════════════════════
RESEARCH DETAILS:
═══════════════════════════════════════
Title/Topic: {topik}
Independent Variable(s): {variabel_bebas}
Dependent Variable(s): {variabel_terikat}
Subjects/Population: {subjek}
Analysis Method: {metode_analisis}
Duration/Period: {durasi}

Data Summary / Findings:
{ringkasan_data if ringkasan_data else "Not provided — generate based on typical findings for this type of research"}

Key References:
{referensi_kunci if referensi_kunci else "Generate appropriate references based on the research topic (APA 7th edition)"}

═══════════════════════════════════════
WRITING INSTRUCTIONS:
═══════════════════════════════════════
1. Do NOT write shallow or generic content
2. Every argument must be grounded in the data provided
3. Use proper academic hedging language
4. Include statistical values where relevant (mean, SD, p-value, effect size)
5. Connect findings to existing literature critically
6. Follow IMRAD structure strictly
7. Format each section with clear headers using ## for main sections
8. For Full Paper: generate all sections completely, do not truncate

Generate the article now:"""
    
    return prompt

def create_docx(content, topik, target, bahasa):
    doc = Document()
    
    # Title
    title = doc.add_heading(topik or "Sports Science Article", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Target: {target} | Language: {bahasa} | Generated by SSPE v1.0").italic = True
    
    doc.add_paragraph("─" * 80)
    
    # Content
    lines = content.split('\n')
    for line in lines:
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        elif line.strip():
            doc.add_paragraph(line)
    
    # Footer
    doc.add_paragraph("─" * 80)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Generated by Sports Science Publication Engine (SSPE) · Prof Bayu").italic = True
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── Generate button ───────────────────────────────────────────────────────────
st.markdown("---")
col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
with col_btn2:
    generate_btn = st.button("🚀 GENERATE ARTIKEL ILMIAH", use_container_width=True)

# ── Output ────────────────────────────────────────────────────────────────────
if generate_btn:
    if not api_key:
        st.error("⚠️ Masukkan API Key Anthropic di sidebar kiri!")
    elif not topik:
        st.error("⚠️ Topik penelitian wajib diisi!")
    else:
        st.markdown("## 📄 Hasil Generate Artikel")
        
        prompt = build_prompt(
            jenis, target, bagian, bahasa, topik,
            variabel_bebas, variabel_terikat, subjek,
            metode_analisis, durasi, ringkasan_data, referensi_kunci
        )
        
        try:
            client = anthropic.Anthropic(api_key=api_key)
            
            with st.spinner("⏳ AI sedang menulis artikel... (estimasi 30-60 detik)"):
                output_placeholder = st.empty()
                full_response = ""
                
                with client.messages.stream(
                    model="claude-sonnet-4-6",
                    max_tokens=8000,
                    messages=[{"role": "user", "content": prompt}]
                ) as stream:
                    for text in stream.text_stream:
                        full_response += text
                        output_placeholder.markdown(
                            f'<div class="output-box">{full_response}▌</div>',
                            unsafe_allow_html=True
                        )
                
                output_placeholder.markdown(
                    f'<div class="output-box">{full_response}</div>',
                    unsafe_allow_html=True
                )
            
            st.success("✅ Artikel berhasil digenerate!")
            
            # Download button
            docx_buf = create_docx(full_response, topik, target, bahasa)
            st.download_button(
                label="📥 Download Artikel (.docx)",
                data=docx_buf,
                file_name=f"SSPE_{topik[:30].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            # Token usage info
            st.info(f"💡 Tip: Simpan hasil ini sebelum generate ulang.")
            
        except anthropic.AuthenticationError:
            st.error("❌ API Key tidak valid. Periksa kembali API Key Anda.")
        except anthropic.RateLimitError:
            st.error("❌ Rate limit tercapai. Tunggu beberapa menit lalu coba lagi.")
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

