"""
citations.py — Modul sitasi untuk Research Co-Pilot (Prof. Bayu)

Fungsi modul ini:
  1. Menormalkan data referensi dari OpenAlex / Crossref / Semantic Scholar
  2. Menyusun Daftar Pustaka APA 7 dengan DOI eksplisit
  3. Membuat sitasi dalam teks BISA DIKLIK -> menuju https://doi.org/...
     - versi HTML  (untuk pratinjau di layar Streamlit)
     - versi .docx (hyperlink asli Word, warna bisa diatur)
  4. Ekspor RIS (Zotero / Mendeley) dan BibTeX (LaTeX / Overleaf)

ATURAN JUJUR YANG DITEGAKKAN DI SINI:
  - Tidak ada DOI  ->  TIDAK ditautkan. Tidak pernah mengarang URL.
  - Sitasi yang tidak cocok dengan referensi terverifikasi -> dibiarkan
    polos (tidak jadi tautan), dan bisa dilaporkan lewat audit_citations().

File ini berdiri sendiri. Tidak mengubah apa pun di sspe_app.py.
"""

import re
import html as _html

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ==========================================================================
# 1. NORMALISASI REFERENSI
# ==========================================================================

def _clean_doi(doi):
    """Ambil DOI murni (10.xxxx/yyy) dari bentuk apa pun. None kalau tidak sah."""
    if not doi:
        return None
    d = str(doi).strip()
    d = re.sub(r'^https?://(dx\.)?doi\.org/', '', d, flags=re.I)
    d = re.sub(r'^doi:\s*', '', d, flags=re.I)
    d = d.strip().rstrip('.')
    # DOI yang sah selalu diawali "10." dan mengandung "/"
    if not d.lower().startswith('10.') or '/' not in d:
        return None
    return d


def doi_url(ref):
    """URL DOI lengkap, atau None kalau referensi tidak punya DOI sah."""
    d = _clean_doi(ref.get('doi'))
    return f"https://doi.org/{d}" if d else None


def _parse_authors(raw):
    """
    Terima berbagai bentuk penulis dan kembalikan list dict:
        [{'family': 'Suherman', 'given': 'Ayi'}, ...]

    Bentuk yang diterima:
      - ["Ayi Suherman", "Budi Santoso"]                     (OpenAlex display_name)
      - [{"family": "Suherman", "given": "Ayi"}]             (Crossref)
      - [{"name": "Ayi Suherman"}] / [{"display_name": ...}] (Semantic Scholar / OpenAlex)
      - "Suherman, A." (satu string)
    """
    if not raw:
        return []
    if isinstance(raw, str):
        raw = [raw]

    out = []
    for a in raw:
        if isinstance(a, dict):
            fam = (a.get('family') or a.get('lastName') or '').strip()
            giv = (a.get('given') or a.get('firstName') or '').strip()
            if fam:
                out.append({'family': fam, 'given': giv})
                continue
            nama = (a.get('name') or a.get('display_name')
                    or a.get('author') or '').strip()
        else:
            nama = str(a).strip()

        if not nama:
            continue

        if ',' in nama:                      # "Suherman, Ayi"
            fam, giv = nama.split(',', 1)
            out.append({'family': fam.strip(), 'given': giv.strip()})
        else:                                # "Ayi Suherman" -> kata terakhir = marga
            bagian = nama.split()
            if len(bagian) == 1:
                out.append({'family': bagian[0], 'given': ''})
            else:
                out.append({'family': bagian[-1], 'given': ' '.join(bagian[:-1])})
    return out


def normalize_ref(raw):
    """
    Ubah satu referensi mentah menjadi bentuk baku yang dipakai modul ini.
    Toleran terhadap beda nama kunci antar sumber.
    """
    g = raw.get

    tahun = g('year') or g('publication_year') or g('published_year')
    try:
        tahun = int(str(tahun)[:4])
    except (TypeError, ValueError):
        tahun = None

    return {
        '_norm': True,
        'authors': _parse_authors(g('authors') or g('author') or g('penulis')),
        'year':    tahun,
        'title':   (g('title') or g('judul') or '').strip().rstrip('.'),
        'journal': (g('journal') or g('venue') or g('container_title')
                    or g('host_venue') or g('source') or '').strip(),
        'volume':  str(g('volume') or '').strip(),
        'issue':   str(g('issue') or '').strip(),
        'pages':   str(g('pages') or g('page') or '').strip(),
        'doi':     _clean_doi(g('doi')),
        'abstract': (g('abstract') or '').strip(),
        'url':     (g('url') or '').strip(),
    }


# ==========================================================================
# 2. DAFTAR PUSTAKA — APA 7
# ==========================================================================

def _inisial(given):
    """'Ayi Rahmat' -> 'A. R.'"""
    if not given:
        return ''
    return ' '.join(f"{p[0].upper()}." for p in given.replace('.', ' ').split() if p)


def _authors_apa(authors):
    """Blok penulis APA 7 untuk Daftar Pustaka."""
    if not authors:
        return ''
    potong = []
    for a in authors:
        ini = _inisial(a['given'])
        potong.append(f"{a['family']}, {ini}" if ini else a['family'])

    if len(potong) == 1:
        return potong[0]
    if len(potong) <= 20:
        return ', '.join(potong[:-1]) + ', & ' + potong[-1]
    # APA 7: >20 penulis -> 19 pertama, elipsis, penulis terakhir
    return ', '.join(potong[:19]) + ', . . . ' + potong[-1]


def format_apa(ref):
    """Satu entri Daftar Pustaka APA 7, lengkap dengan DOI bila ada."""
    r = ref if ref.get('_norm') else normalize_ref(ref)

    bagian = []
    penulis = _authors_apa(r['authors'])
    bagian.append(penulis if penulis else '(Tanpa penulis)')
    bagian.append(f"({r['year']})." if r['year'] else '(n.d.).')

    if r['title']:
        bagian.append(f"{r['title']}.")

    if r['journal']:
        jur = r['journal']
        if r['volume']:
            jur += f", {r['volume']}"
            if r['issue']:
                jur += f"({r['issue']})"
        if r['pages']:
            jur += f", {r['pages']}"
        bagian.append(jur + '.')

    tautan = doi_url(r)
    if tautan:
        bagian.append(tautan)
    elif r['url']:
        bagian.append(r['url'])

    return ' '.join(bagian)


def intext_apa(ref):
    """Sitasi dalam teks yang seharusnya, mis. '(Suherman et al., 2021)'."""
    r = ref if ref.get('_norm') else normalize_ref(ref)
    a = r['authors']
    thn = r['year'] or 'n.d.'
    if not a:
        return f"(Anonim, {thn})"
    if len(a) == 1:
        return f"({a[0]['family']}, {thn})"
    if len(a) == 2:
        return f"({a[0]['family']} & {a[1]['family']}, {thn})"
    return f"({a[0]['family']} et al., {thn})"


# ==========================================================================
# 3. INDEKS & PENCOCOKAN SITASI
# ==========================================================================

# (Marga, 2021) / (Marga & Lain, 2021) / (Marga et al., 2021)
_POLA_KURUNG = re.compile(r'\(([^()]{3,240}?)\)')
# Marga (2021) — sitasi naratif
_POLA_NARATIF = re.compile(
    r"\b([A-ZÀ-Þ][\w'’\-]+(?:\s+(?:et\s+al\.|&\s+[A-ZÀ-Þ][\w'’\-]+|dan\s+[A-ZÀ-Þ][\w'’\-]+))?)"
    r"\s+\((\d{4})\)"
)


def build_index(refs):
    """
    Bangun indeks pencarian dari daftar referensi TERVERIFIKASI.
    Kunci: (marga_penulis_pertama_lowercase, tahun) -> referensi ternormalisasi.
    Hanya referensi ber-DOI yang akan bisa ditautkan.
    """
    idx = {}
    for raw in refs or []:
        r = normalize_ref(raw)
        if not r['authors'] or not r['year']:
            continue
        kunci = (r['authors'][0]['family'].lower(), r['year'])
        idx.setdefault(kunci, r)
    return idx


def _cocokkan(potongan, idx):
    """
    Cocokkan satu potongan sitasi (mis. 'Suherman et al., 2021') ke indeks.
    Kembalikan referensi, atau None kalau tidak cocok.
    """
    m = re.search(r'\b(1[5-9]\d{2}|20\d{2})\b', potongan)
    if not m:
        return None
    tahun = int(m.group(1))

    teks_nama = potongan[:m.start()]
    teks_nama = re.sub(r'\bet\s+al\.?', ' ', teks_nama, flags=re.I)
    teks_nama = re.sub(r'[,;&]|(\bdan\b)', ' ', teks_nama, flags=re.I)
    kata = [k for k in teks_nama.split() if len(k) > 1]
    if not kata:
        return None

    # coba marga pertama, lalu marga terakhir (jaga-jaga urutan nama Indonesia)
    for kandidat in (kata[0], kata[-1]):
        r = idx.get((kandidat.lower().strip(".'’-"), tahun))
        if r:
            return r
    return None


def audit_citations(teks, refs):
    """
    Periksa teks bab: sitasi mana yang cocok dengan referensi terverifikasi,
    dan mana yang menggantung. Dipakai untuk peringatan di layar.

    Kembali: {'cocok': [...], 'menggantung': [...], 'tanpa_doi': [...]}
    """
    idx = build_index(refs)
    cocok, menggantung, tanpa_doi = [], [], []

    def _catat(label, r):
        if r is None:
            menggantung.append(label)
        elif doi_url(r):
            cocok.append(label)
        else:
            tanpa_doi.append(label)

    for m in _POLA_KURUNG.finditer(teks or ''):
        isi = m.group(1)
        if not re.search(r'\b(1[5-9]\d{2}|20\d{2})\b', isi):
            continue
        if re.fullmatch(r'\s*\d{4}[a-z]?\s*', isi):
            continue          # ini bagian dari sitasi naratif: Metzler (2017)
        for potongan in isi.split(';'):
            _catat(potongan.strip(), _cocokkan(potongan, idx))

    for m in _POLA_NARATIF.finditer(teks or ''):
        _catat(f"{m.group(1)} ({m.group(2)})",
               _cocokkan(f"{m.group(1)}, {m.group(2)}", idx))

    return {'cocok': cocok, 'menggantung': menggantung, 'tanpa_doi': tanpa_doi}


# ==========================================================================
# 4. PEMOTONGAN TEKS JADI POTONGAN BERTAUTAN
# ==========================================================================

def _segments(teks, idx):
    """
    Pecah teks jadi urutan (potongan_teks, url_atau_None).
    Inilah inti fitur "sitasi bisa diklik" — dipakai HTML maupun .docx.
    """
    tanda = []   # (mulai, selesai, url)

    for m in _POLA_KURUNG.finditer(teks):
        isi, awal = m.group(1), m.start(1)
        if not re.search(r'\b(1[5-9]\d{2}|20\d{2})\b', isi):
            continue
        if re.fullmatch(r'\s*\d{4}[a-z]?\s*', isi):
            continue          # ini bagian dari sitasi naratif: Metzler (2017)
        geser = 0
        for potongan in isi.split(';'):
            r = _cocokkan(potongan, idx)
            u = doi_url(r) if r else None
            if u:
                kiri = len(potongan) - len(potongan.lstrip())
                kanan = len(potongan) - len(potongan.rstrip())
                tanda.append((awal + geser + kiri,
                              awal + geser + len(potongan) - kanan, u))
            geser += len(potongan) + 1

    for m in _POLA_NARATIF.finditer(teks):
        r = _cocokkan(f"{m.group(1)}, {m.group(2)}", idx)
        u = doi_url(r) if r else None
        if u and not any(a <= m.start() < b for a, b, _ in tanda):
            tanda.append((m.start(), m.end(), u))

    tanda.sort()
    hasil, kursor = [], 0
    for a, b, u in tanda:
        if a < kursor:
            continue
        if a > kursor:
            hasil.append((teks[kursor:a], None))
        hasil.append((teks[a:b], u))
        kursor = b
    if kursor < len(teks):
        hasil.append((teks[kursor:], None))
    return hasil


# ==========================================================================
# 5. PRATINJAU DI LAYAR (HTML untuk Streamlit)
# ==========================================================================

def linkify_html(teks, refs, warna='#E8730A'):
    """
    Versi HTML: sitasi jadi tautan berwarna (default oranye, sesuai tema).
    Pakai di Streamlit dengan st.markdown(..., unsafe_allow_html=True).
    """
    idx = build_index(refs)
    keluar = []
    for potongan, u in _segments(teks or '', idx):
        aman = _html.escape(potongan)
        if u:
            keluar.append(
                f'<a href="{_html.escape(u, quote=True)}" target="_blank" '
                f'style="color:{warna};text-decoration:none;'
                f'border-bottom:1px dotted {warna};">{aman}</a>'
            )
        else:
            keluar.append(aman)
    return ''.join(keluar)


# ==========================================================================
# 6. HYPERLINK ASLI DI WORD (.docx)
# ==========================================================================

def add_hyperlink(paragraph, teks, url, biru=False):
    """
    Sisipkan hyperlink Word yang benar-benar aktif.
    python-docx tidak punya fungsi ini, jadi ditulis manual lewat XML.

    biru=False -> tampil hitam polos, TAPI tautan tetap hidup (Ctrl+klik).
                  Ini default, supaya format cetak tetap sesuai APA 7.
    """
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    tautan = OxmlElement('w:hyperlink')
    tautan.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if biru:
        c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
        u = OxmlElement('w:u');     u.set(qn('w:val'), 'single'); rPr.append(u)
    run.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = teks
    run.append(t)

    tautan.append(run)
    paragraph._p.append(tautan)
    return tautan


def add_paragraph_with_links(doc, teks, refs_atau_idx, biru=False, style=None):
    """
    Tambahkan satu paragraf ke dokumen Word dengan sitasi yang bisa diklik.
    Ganti pemakaian doc.add_paragraph(teks) dengan fungsi ini.
    """
    idx = refs_atau_idx if isinstance(refs_atau_idx, dict) \
        else build_index(refs_atau_idx)
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for potongan, u in _segments(teks or '', idx):
        if u:
            add_hyperlink(p, potongan, u, biru=biru)
        else:
            p.add_run(potongan)
    return p


def add_bibliography(doc, refs, biru=False, judul='DAFTAR PUSTAKA'):
    """Tulis Daftar Pustaka APA 7; DOI ditulis penuh dan bisa diklik."""
    if judul:
        doc.add_heading(judul, level=1)

    entri = sorted(
        [normalize_ref(r) for r in (refs or [])],
        key=lambda r: ((r['authors'][0]['family'] if r['authors'] else 'zzz').lower(),
                       r['year'] or 9999)
    )
    for r in entri:
        teks = format_apa(r)
        tautan = doi_url(r)
        p = doc.add_paragraph()
        if tautan and teks.endswith(tautan):
            p.add_run(teks[:-len(tautan)])
            add_hyperlink(p, tautan, tautan, biru=biru)
        else:
            p.add_run(teks)
    return doc


# ==========================================================================
# 7. EKSPOR RIS & BIBTEX
# ==========================================================================

def to_ris(refs):
    """RIS untuk diimpor ke Zotero / Mendeley."""
    baris = []
    for raw in refs or []:
        r = normalize_ref(raw)
        baris.append('TY  - JOUR')
        for a in r['authors']:
            ini = _inisial(a['given'])
            baris.append(f"AU  - {a['family']}, {ini}" if ini
                         else f"AU  - {a['family']}")
        if r['year']:    baris.append(f"PY  - {r['year']}")
        if r['title']:   baris.append(f"TI  - {r['title']}")
        if r['journal']: baris.append(f"JO  - {r['journal']}")
        if r['volume']:  baris.append(f"VL  - {r['volume']}")
        if r['issue']:   baris.append(f"IS  - {r['issue']}")
        if r['pages']:
            hal = r['pages'].replace('--', '-')
            if '-' in hal:
                sp, ep = hal.split('-', 1)
                baris.append(f"SP  - {sp.strip()}")
                baris.append(f"EP  - {ep.strip()}")
            else:
                baris.append(f"SP  - {hal}")
        if r['doi']:
            baris.append(f"DO  - {r['doi']}")
            baris.append(f"UR  - https://doi.org/{r['doi']}")
        elif r['url']:
            baris.append(f"UR  - {r['url']}")
        if r['abstract']:
            baris.append(f"AB  - {r['abstract']}")
        baris.append('ER  - ')
        baris.append('')
    return '\r\n'.join(baris)


def _kunci_bibtex(r, dipakai):
    marga = re.sub(r'[^A-Za-z]', '', r['authors'][0]['family']) if r['authors'] else 'anon'
    dasar = f"{marga.lower()}{r['year'] or 'nd'}"
    kunci, n = dasar, 1
    while kunci in dipakai:
        n += 1
        kunci = f"{dasar}{chr(96 + n)}"
    dipakai.add(kunci)
    return kunci


def to_bibtex(refs):
    """BibTeX untuk pengguna LaTeX / Overleaf."""
    dipakai, keluar = set(), []
    for raw in refs or []:
        r = normalize_ref(raw)
        f = []
        if r['authors']:
            f.append(('author', ' and '.join(
                f"{a['family']}, {a['given']}" if a['given'] else a['family']
                for a in r['authors'])))
        if r['title']:   f.append(('title', r['title']))
        if r['journal']: f.append(('journal', r['journal']))
        if r['year']:    f.append(('year', str(r['year'])))
        if r['volume']:  f.append(('volume', r['volume']))
        if r['issue']:   f.append(('number', r['issue']))
        if r['pages']:   f.append(('pages', r['pages'].replace('-', '--')))
        if r['doi']:
            f.append(('doi', r['doi']))
            f.append(('url', f"https://doi.org/{r['doi']}"))
        isi = ',\n'.join(f"  {k} = {{{v}}}" for k, v in f)
        keluar.append(f"@article{{{_kunci_bibtex(r, dipakai)},\n{isi}\n}}")
    return '\n\n'.join(keluar)
